import os
import io
import docx
import logging
import requests

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

from seadoc_converter.config import SEAHUB_SERVICE_URL
from seadoc_converter.converter.utils import gen_jwt_auth_header

logger = logging.getLogger(__name__)

DEFAULT_CALLOUT_COLOR = 'fef7e0'


def get_image_content_url(file_uuid, image_name):

    payload = {'file_uuid': file_uuid}

    url = f'{SEAHUB_SERVICE_URL}/api/v2.1/seadoc/image-download-link/{file_uuid}/'
    params = {'image_name': image_name}
    headers = gen_jwt_auth_header(payload)

    resp = requests.get(url, params, headers=headers)
    if resp.status_code == 200:
        return resp.json().get('download_link')
    else:
        logger.error(resp.__dict__)
        return ""


def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def sdoc2docx(file_content_json, file_uuid, username):

    def add_hyperlink(paragraph, url, text, color):
        """
        A function that places a hyperlink within a paragraph object.

        :param paragraph: The paragraph we are adding the hyperlink to.
        :param url: A string containing the required url
        :param text: The text displayed for the url
        :return: The hyperlink object
        """

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Add color if it is given
        if color:
            c = docx.oxml.shared.OxmlElement('w:color')
            c.set(docx.oxml.shared.qn('w:val'), color)
            rPr.append(c)

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink

    def extract_text_in_table_recursively(data):

        text_list = []
        if isinstance(data, list):
            for item in data:
                text_list.extend(extract_text_in_table_recursively(item))
        elif isinstance(data, dict):
            if 'text' in data:
                text_list.append(data['text'])
            for key, value in data.items():
                text_list.extend(extract_text_in_table_recursively(value))

        return text_list

    def search_sdoc_node_recursively(children_list, type_sq=[], top_type='', top_style=''):

        if 'text' in children_list[0]:
            if top_type == "ordered_list" and type_sq.count('ordered_list') == 1:
                type_content_list.append(['ordered_list_2', children_list, top_style])
            elif top_type == "ordered_list" and type_sq.count('ordered_list') >= 2:
                type_content_list.append(['ordered_list_3', children_list, top_style])
            elif top_type == "unordered_list" and type_sq.count('unordered_list') == 1:
                type_content_list.append(['ordered_list_2', children_list, top_style])
            elif top_type == "unordered_list" and type_sq.count('unordered_list') >= 2:
                type_content_list.append(['ordered_list_3', children_list, top_style])
            else:
                type_content_list.append([top_type, children_list, top_style])
        else:
            if top_type == 'table':
                table_text_list = extract_text_in_table_recursively(children_list)
                sub_length = len(children_list[0]['children'])
                new_table_text_list = []
                for i in range(0, len(table_text_list), sub_length):
                    new_table_text_list.append(table_text_list[i:i + sub_length])

                type_content_list.append([top_type, new_table_text_list, top_style])
            else:
                for children in children_list:
                    current_type = children.get('type', 'no type')
                    sub_children_list = children.get('children', [])
                    if sub_children_list:
                        search_sdoc_node_recursively(sub_children_list,
                                                 type_sq + [current_type],
                                                 top_type=top_type, top_style=top_style)

    sdoc_node_list = file_content_json.get('elements', [])
    type_content_list = []
    for sdoc_node in sdoc_node_list:
        top_sdoc_type = sdoc_node.get('type', '')
        children_list = sdoc_node.get('children', [])
        style = sdoc_node.get('style', '')
        if children_list:
            search_sdoc_node_recursively(children_list, top_type=top_sdoc_type, top_style=style)

    document = Document()

    for type_content in type_content_list:

        sdoc_type = type_content[0]
        content = type_content[1]
        style = type_content[2]

        if sdoc_type == 'title':
            docx_paragraph = document.add_heading(level=0)
        if sdoc_type == 'subtitle':
            docx_paragraph = document.add_paragraph(style="Subtitle")
        if sdoc_type == 'header1':
            docx_paragraph = document.add_heading(level=1)
        if sdoc_type == 'header2':
            docx_paragraph = document.add_heading(level=2)
        if sdoc_type == 'header3':
            docx_paragraph = document.add_heading(level=3)
        if sdoc_type == 'header4':
            docx_paragraph = document.add_heading(level=4)
        if sdoc_type == 'header5':
            docx_paragraph = document.add_heading(level=5)
        if sdoc_type == 'header6':
            docx_paragraph = document.add_heading(level=6)
        if sdoc_type == 'paragraph':
            docx_paragraph = document.add_paragraph()
        if sdoc_type == 'blockquote':
            docx_paragraph = document.add_paragraph(style="Intense Quote")
        if sdoc_type == 'ordered_list':
            docx_paragraph = document.add_paragraph(style="List Number")
        if sdoc_type == 'ordered_list_2':
            docx_paragraph = document.add_paragraph(style="List Number 2")
        if sdoc_type == 'ordered_list_3':
            docx_paragraph = document.add_paragraph(style="List Number 3")
        if sdoc_type in ('unordered_list', 'check_list_item'):
            docx_paragraph = document.add_paragraph(style="List Bullet")
        if sdoc_type == 'unordered_list_2':
            docx_paragraph = document.add_paragraph(style="List Bullet 2")
        if sdoc_type == 'unordered_list_3':
            docx_paragraph = document.add_paragraph(style="List Bullet 3")

        if sdoc_type == 'code_block':

            docx_paragraph = document.add_paragraph(style="No Spacing")
            docx_paragraph.paragraph_format.left_indent = Inches(0.2)

            for text_dict in content:
                text = text_dict.get('text', '')
                run = docx_paragraph.add_run(text)
                run.font.size = Pt(10)
                run.font.name = 'Courier New'

        elif sdoc_type == 'paragraph' and \
                any(item.get('type') == 'link' for item in content):

            # add hyperlink to docx

            # ['paragraph',
            #   [{'id': 'TQdHtyxhQfm8ipm76cVKKg', 'text': ''},
            #    {'children': [{'id': 'VFGENWpbTNeMRb-16QgdNA',
            #                   'text': '127.0.0.1 link title'}],
            #     'href': 'http://127.0.0.1/link-address/',
            #     'id': 'Co9L-c-SQmWk4yxHSXu5tg',
            #     'title': '127.0.0.1 link title',
            #     'type': 'link'},
            #    {'id': 'Pwqf3nbSTWmIFbwrFo1Eow', 'text': ''}]],

            link_href = ''
            link_title = ''
            for item in content:
                if 'href' in item:
                    link_href = item['href']
                if 'title' in item:
                    link_title = item['title']

            docx_paragraph = document.add_paragraph()
            add_hyperlink(docx_paragraph, link_href, link_title, "0000FF")

        elif sdoc_type == 'paragraph' and \
                any(item.get('type') in ('sdoc_link', 'file_link') for item in content):

            # add sdoc/file link to docx

            # ['paragraph',
            #  [{'id': 'D8omdcCLR4eLB3o4f0yOxw', 'text': ' '},
            #   {'children': [{'id': 'KFM5z7zvTaOcZyaT1zBhHQ', 'text': '987.sdoc'}],
            #    'display_type': 'icon_link',
            #    'doc_uuid': '45b266e4-17a5-475d-b601-10aa8001ea80',
            #    'id': 'bIwxx0mMQVKRFo3LlYwf6A',
            #    'title': '987.sdoc',
            #    'type': 'sdoc_link'},
            #   {'id': 'G5WmlQ4tSpO4IH5CDFCdUA', 'text': ' '}]],

            doc_uuid = ''
            doc_title = ''
            for item in content:
                if 'doc_uuid' in item:
                    doc_uuid = item['doc_uuid']
                if 'title' in item:
                    doc_title = item['title']

            doc_url = f'{SEAHUB_SERVICE_URL}/api/v2.1/seadoc/file/{doc_uuid}/'
            docx_paragraph = document.add_paragraph()
            add_hyperlink(docx_paragraph, doc_url, doc_title, "0000FF")

        elif sdoc_type in ('paragraph', 'image_block') and \
                any(item.get('type') == 'image' for item in content):

            # add image to docx

            # ['paragraph',
            #  [{'id': 'VL579VQRQdOjJCKkjRXXNA', 'text': ''},
            #   {'children': [{'id': 'dp7gIr5aSEa6GtK3-vi68g', 'text': ''}],
            #    'data': {'src': '/image-1702627227876.png'},
            #    'id': 'TEPevi-FQo-unZRBSlnd3A',
            #    'type': 'image'},
            #   {'id': 'SQjLfnvBSimn695OZtyGnw', 'text': ''}]],

            image_file_path = ''
            for item in content:
                if 'data' in item:
                    image_file_path = item['data']['src']

            image_name = os.path.basename(image_file_path)
            image_content_url = get_image_content_url(file_uuid, image_name)
            if image_content_url:
                resp = requests.get(image_content_url)
                image_content = resp.content
                try:
                    document.add_picture(io.BytesIO(image_content), width=Inches(5))
                except Exception as e:
                    logger.debug('add image to docx failed: file_uuid: %s, image_path: %s, error: %s', file_uuid, image_content_url, e)
            else:
                logger.exception(f'can not get image content: {file_uuid} {image_file_path}')

        elif sdoc_type == 'table':

            # add table to docx

            # ['table', [['1', '2', '3', '4'], ['a', 'b', 'c', 'd']]]

            table = document.add_table(rows=len(content), cols=len(content[0]))

            def fulfill_table(table, content):
                for i, row in enumerate(content):
                    for j, value in enumerate(row):
                        table.cell(i, j).text = value

            fulfill_table(table, content)

        elif sdoc_type == 'callout':

            docx_paragraph = document.add_paragraph()

            # Create XML element
            shd = OxmlElement('w:shd')
            color = DEFAULT_CALLOUT_COLOR
            if style:
                color = style.get('background_color').lstrip('#')

            # Add attributes to the element
            shd.set(qn('w:fill'), color)

            # Make sure the paragraph styling element exists
            docx_paragraph.paragraph_format.element.get_or_add_pPr()

            # Append the shading element
            docx_paragraph.paragraph_format.element.pPr.append(shd)

            for text_dict in content:
                text = text_dict.get('text', '')
                run = docx_paragraph.add_run(text)

                bold = text_dict.get('bold', False)
                run.bold = True if bold else False
                italic = text_dict.get('italic', False)
                run.italic = True if italic else False
                if hex_color := text_dict.get('color', None):
                    rgb_color = hex_to_rgb(hex_color)
                    run.font.color.rgb = RGBColor(*rgb_color)
                if font_name := text_dict.get('font', None):
                    run.font.name = font_name
                if font_size := text_dict.get('font_size', None):  
                    run.font.size = Pt(font_size)
        else:

            for text_dict in content:

                text = text_dict.get('text', '') or text_dict.get('href', '')
                run = docx_paragraph.add_run(text)

                bold = text_dict.get('bold', False)
                run.bold = True if bold else False

                italic = text_dict.get('italic', False)
                run.italic = True if italic else False
                if hex_color := text_dict.get('color', None):
                    rgb_color = hex_to_rgb(hex_color)
                    run.font.color.rgb = RGBColor(*rgb_color)
                if font_name := text_dict.get('font', None):
                    run.font.name = font_name
                if font_size := text_dict.get('font_size', None):  
                    run.font.size = Pt(font_size)

    memory_stream = io.BytesIO()
    document.save(memory_stream)
    docx_content = memory_stream.getvalue()
    return docx_content
